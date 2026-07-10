"""Routers that deliver a submitted memo to Notesnook or Google Drive."""
import html
import re
import shutil
from datetime import datetime
from pathlib import Path

import requests


# A note is stored as plain text, so a list is just its Markdown line — which is
# what the editor writes and reads back. Both destinations turn those lines into
# real lists: HTML for Notesnook, Word's list styles for a Drive .docx.
_BULLET = re.compile(r"^\s*[-*•]\s+(.*)$")
_NUMBER = re.compile(r"^\s*\d+[.)]\s+(.*)$")


def _list_item(line):
    """The item text and its list tag ("ul"/"ol"), or (None, None) for prose."""
    bullet = _BULLET.match(line)
    if bullet:
        return bullet.group(1).strip(), "ul"
    number = _NUMBER.match(line)
    if number:
        return number.group(1).strip(), "ol"
    return None, None


def _text_to_html(text):
    parts = []
    open_tag = None
    for line in text.split("\n"):
        item, tag = _list_item(line)
        if tag != open_tag:
            if open_tag:
                parts.append(f"</{open_tag}>")
            if tag:
                parts.append(f"<{tag}>")
            open_tag = tag
        if tag:
            parts.append(f"<li>{html.escape(item)}</li>")
        elif line.strip():
            parts.append(f"<p>{html.escape(line.strip())}</p>")
    if open_tag:
        parts.append(f"</{open_tag}>")
    return "".join(parts) or "<p></p>"


def _default_title(timestamp):
    """Name an unnamed memo the way Notesnook names untitled notes ("Note $date$
    $time$"), from an ISO timestamp. Notesnook's Inbox API requires a non-empty
    title, so this always returns one even when the timestamp is missing.

    The time is to the second, not the minute: two unnamed memos recorded in the
    same minute would otherwise share a title, and same-titled notes collapse to one
    in the inbox — silently dropping every second recording made within a minute."""
    try:
        made = datetime.fromisoformat(timestamp)
    except (TypeError, ValueError):
        return "Voice note"
    hour = made.hour % 12 or 12
    meridiem = "AM" if made.hour < 12 else "PM"
    return f"Note {made:%Y-%m-%d} {hour}:{made:%M}:{made:%S} {meridiem}"


class NotesnookRouter:
    """Create a note via the Notesnook Inbox API (POST https://inbox.notesnook.com/)."""

    ENDPOINT = "https://inbox.notesnook.com/"

    def __init__(self, api_key, *, source="highdeas", post=requests.post):
        self._api_key = api_key
        self._source = source
        self._post = post

    def route(self, memo):
        response = self._post(
            self.ENDPOINT,
            headers={"Authorization": self._api_key, "Content-Type": "application/json"},
            json={
                "title": memo.name or _default_title(memo.recorded_at or memo.created_at),
                "type": "note",
                "source": self._source,
                "version": 1,
                "content": {"type": "html", "data": _text_to_html(memo.transcript)},
            },
            timeout=30,
        )
        response.raise_for_status()


class Router:
    """Dispatch a memo to the Notesnook or Drive router based on its chosen route."""

    def __init__(self, notesnook, drive=None):
        self._notesnook = notesnook
        self._drive = drive

    def __call__(self, memo):
        if memo.route == "drive":
            if self._drive is not None:
                self._drive.route(memo)
        else:
            self._notesnook.route(memo)


def _today():
    return datetime.now().strftime("%Y_%m_%d")


def _sanitize_filename(name):
    cleaned = re.sub(r'[<>:"/\\|?*]', "", name).strip()
    return cleaned or "untitled"


_LIST_STYLE = {"ul": "List Bullet", "ol": "List Number"}


def write_docx(path, text):
    from docx import Document

    document = Document()
    for line in text.split("\n"):
        item, tag = _list_item(line)
        if tag:
            document.add_paragraph(item, style=_LIST_STYLE[tag])
        else:
            document.add_paragraph(line)
    document.save(str(path))


class DriveMusicRouter:
    """Copy a music memo into a dated folder under the Drive base, with an optional doc.

    The original stays in the inbox so the service can also retire it to the local
    bin — the memo is then recoverable there for 90 days regardless of what happens
    to the Drive copy."""

    def __init__(self, inbox_dir, drive_base, *, today=_today, write_doc=write_docx, copy=shutil.copy2):
        self._inbox = Path(inbox_dir)
        self._base = Path(drive_base)
        self._today = today
        self._write_doc = write_doc
        self._copy = copy

    def route(self, memo):
        folder = self._base / f"_{self._today()}_NOT_YET_PROCESSED_MUSIC"
        folder.mkdir(parents=True, exist_ok=True)
        source = self._inbox / memo.audio_filename
        base = _sanitize_filename(memo.name or Path(memo.audio_filename).stem)
        self._copy(str(source), str(folder / (base + source.suffix)))
        if memo.transcript.strip():
            self._write_doc(folder / (base + ".docx"), memo.transcript)
